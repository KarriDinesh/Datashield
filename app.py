import re
import os
import io
import uuid
import copy
from flask import Flask, request, render_template, flash, redirect, url_for, send_file, session


# 1. PDF Support
try:
    from pypdf import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# 2. Word (.docx) Support
try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# 3. Excel (.xlsx) Support
try:
    import openpyxl
    EXCEL_SUPPORT = True
except ImportError:
    EXCEL_SUPPORT = False

app = Flask(__name__)
app.secret_key = 'Masker'

PROCESSED_FILES = {}
UPLOAD_CACHE = {} # Cache to hold uploaded files during the "Analyze" step


PATTERNS = {
    'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
    'phone': r'\(?\b[2-9][0-9]{2}\)?[-. ]?[2-9][0-9]{2}[-. ]?[0-9]{4}\b', 
    'credit_card': r'\b(?:\d[ -]*?){13,16}\b',
    'ssn': r'\b\d{3}-\d{2}-\d{4}\b'
}

def get_active_patterns(options):
    """Returns a dict of patterns based on selected options"""
    active = {}
    if options.get('mask_email'): active['email'] = PATTERNS['email']
    if options.get('mask_phone'): active['phone'] = PATTERNS['phone']
    if options.get('mask_credit_card'): active['credit_card'] = PATTERNS['credit_card']
    if options.get('mask_ssn'): active['ssn'] = PATTERNS['ssn']
    return active

def analyze_content(text, options):
    """
    Finds unique matches and returns them for user review.
    Returns: List of dicts [{'type': 'email', 'value': 'bob@ex.com'}, ...]
    """
    findings = []
    seen = set()
    
    active_patterns = get_active_patterns(options)

    for p_type, pattern in active_patterns.items():
        matches = re.finditer(pattern, text)
        for m in matches:
            val = m.group()
            if val not in seen:
                seen.add(val)
                findings.append({'type': p_type, 'value': val})
    
    return findings

def perform_masking_on_text(text, options, ignored_set, stats):
    """Helper function to run regex sub on a string"""
    active_patterns = get_active_patterns(options)
    masked_text = text

    for p_type, pattern in active_patterns.items():
        def replacement_handler(match):
            val = match.group()
            if val in ignored_set:
                return val
            
            # Update stats
            stats[p_type] += 1
            stats['total'] += 1
            
            if p_type == 'email': return "******@***.**"
            elif p_type == 'phone': return "[PHONE-REDACTED]"
            elif p_type == 'credit_card': return "[CC-REDACTED]"
            elif p_type == 'ssn': return "***-**-****"
            return "*****"

        masked_text = re.sub(pattern, replacement_handler, masked_text)
    return masked_text

def mask_content_and_reconstruct_file(text_input, file_bytes, filename, options, ignored_items):
    """
    Masks the content.
    If file_bytes (bytes) is provided and is .docx or .xlsx, it reconstructs the file.
    Otherwise, returns text masking only.
    """
    stats = { 'email': 0, 'phone': 0, 'credit_card': 0, 'ssn': 0, 'total': 0 }
    ignored_set = set(ignored_items)

    # 1. Always generate masked plain text for the UI
    # This populates the 'stats' dictionary that we display in the UI
    sanitized_text_for_ui = perform_masking_on_text(text_input, options, ignored_set, stats)
    
    download_id = None
    
    # 2. Reconstruct File if possible
    if file_bytes and filename:
        ext = filename.lower().split('.')[-1]
        
        output_io = io.BytesIO()
        processed_file_name = f"masked_{filename}"
        is_processed = False
        
        dummy_stats = { 'email': 0, 'phone': 0, 'credit_card': 0, 'ssn': 0, 'total': 0 }
        
        try:
            # Wrap bytes in IO stream
            file_stream = io.BytesIO(file_bytes)

            # --- DOCX ---
            if ext == 'docx' and DOCX_SUPPORT:
                doc = docx.Document(file_stream)
                
                # Mask Paragraphs
                for para in doc.paragraphs:
                    para.text = perform_masking_on_text(para.text, options, ignored_set, dummy_stats)
                
                # Mask Tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell.text = perform_masking_on_text(cell.text, options, ignored_set, dummy_stats)
                
                doc.save(output_io)
                is_processed = True

            # --- XLSX ---
            elif ext == 'xlsx' and EXCEL_SUPPORT:
                wb = openpyxl.load_workbook(file_stream)
                
                for ws in wb.worksheets:
                    for row in ws.iter_rows():
                        for cell in row:
                            if cell.value and isinstance(cell.value, str):
                                cell.value = perform_masking_on_text(cell.value, options, ignored_set, dummy_stats)
                
                wb.save(output_io)
                is_processed = True
            
            # --- PDF / TEXT / OTHERS ---
            else:
                # For others, we just save the masked plain text as a .txt file
                output_io.write(sanitized_text_for_ui.encode('utf-8'))
                processed_file_name = f"masked_{filename}.txt"
                is_processed = True

            if is_processed:
                output_io.seek(0)
                download_id = str(uuid.uuid4())
                PROCESSED_FILES[download_id] = {
                    'data': output_io,
                    'filename': processed_file_name
                }

        except Exception as e:
            print(f"Error reconstruction file: {e}")

    return sanitized_text_for_ui, stats, download_id


@app.route('/download/<download_id>')
def download_file(download_id):
    file_info = PROCESSED_FILES.get(download_id)
    if not file_info:
        flash("File link expired or invalid.", "error")
        return redirect(url_for('index'))
    
    return send_file(
        file_info['data'],
        as_attachment=True,
        download_name=file_info['filename']
    )

@app.route('/', methods=['GET', 'POST'])
def index():
    original_text = ""
    sanitized_text = ""
    detected_items = []
    stats = {}
    download_id = None
    filename = None
    cache_id = None
    
    # Default options
    options = {
        'mask_email': True,
        'mask_phone': True,
        'mask_credit_card': True,
        'mask_ssn': True
    }

    if request.method == 'POST':
        action = request.form.get('action', 'analyze')
        
        # Reload options from form
        options = {
            'mask_email': request.form.get('mask_email') == 'yes',
            'mask_phone': request.form.get('mask_phone') == 'yes',
            'mask_credit_card': request.form.get('mask_credit_card') == 'yes',
            'mask_ssn': request.form.get('mask_ssn') == 'yes'
        }
        
        # Variables to hold file content
        file_bytes = None
        
        if 'file' in request.files and request.files['file'].filename != '':
            file = request.files['file']
            filename = file.filename
            
            # Read bytes into memory for reuse in cache
            file_bytes = file.read()
            
            # Cache the file for the next step (Masking)
            cache_id = str(uuid.uuid4())
            UPLOAD_CACHE[cache_id] = {
                'data': file_bytes,
                'filename': filename
            }

            # Create stream for extractors
            file_stream = io.BytesIO(file_bytes)
            
            # Read text for analysis
            try:
                if filename.lower().endswith('.pdf'):
                    if PDF_SUPPORT:
                        try:
                            reader = PdfReader(file_stream)
                            extracted_pages = []
                            for page in reader.pages:
                                extracted_pages.append(page.extract_text())
                            original_text = "\n".join(extracted_pages)
                        except Exception as e:
                            flash(f"Error parsing PDF: {str(e)}", "error")
                    else:
                        flash("PDF support requires 'pypdf'.", "error")
                
                elif filename.lower().endswith('.docx'):
                    if DOCX_SUPPORT:
                        try:
                            doc = docx.Document(file_stream)
                            full_text = []
                            for para in doc.paragraphs:
                                full_text.append(para.text)
                            original_text = '\n'.join(full_text)
                        except Exception as e:
                            flash(f"Error parsing DOCX: {str(e)}", "error")
                    else:
                         flash("Word support requires 'python-docx'.", "error")

                elif filename.lower().endswith('.xlsx'):
                    if EXCEL_SUPPORT:
                        try:
                            wb = openpyxl.load_workbook(file_stream, data_only=True)
                            full_text = []
                            for sheet in wb.sheetnames:
                                ws = wb[sheet]
                                for row in ws.iter_rows(values_only=True):
                                    row_text = " ".join([str(cell) for cell in row if cell is not None])
                                    full_text.append(row_text)
                            original_text = '\n'.join(full_text)
                        except Exception as e:
                            flash(f"Error parsing Excel: {str(e)}", "error")
                    else:
                        flash("Excel support requires 'openpyxl'.", "error")

                else:
                    # Text/Generic
                    try:
                        original_text = file_bytes.decode('utf-8')
                    except UnicodeDecodeError:
                        try:
                            original_text = file_bytes.decode('latin-1') 
                        except UnicodeDecodeError:
                            original_text = ""
            except Exception as e:
                flash(f"Error reading file: {str(e)}", "error")
        
        else:
            original_text = request.form.get('text_input', '')
            filename = request.form.get('existing_filename', '')
            cache_id = request.form.get('cache_id', '')
            
            # If we are masking and have a cache_id, retrieve the file
            if action == 'mask' and cache_id and cache_id in UPLOAD_CACHE:
                cached = UPLOAD_CACHE[cache_id]
                file_bytes = cached['data']
                # Ensure filename matches (sanity check) or just use cached one
                if not filename: 
                    filename = cached['filename']

        if original_text:
            if action == 'analyze':
                detected_items = analyze_content(original_text, options)
                if not detected_items:
                    flash("No sensitive data found matching selected filters.", "success")
                else:
                    flash("Analysis complete. Select items to ignore, then confirm.", "success")
            
            elif action == 'mask':
                ignored_items = request.form.getlist('ignored_items')
                
                sanitized_text, stats, download_id = mask_content_and_reconstruct_file(
                    original_text, file_bytes, filename, options, ignored_items
                )
                
                if cache_id and cache_id in UPLOAD_CACHE:
                    del UPLOAD_CACHE[cache_id]
                    cache_id = None
                
                flash(f"Masking complete. {stats['total']} items redacted.", "success")

    return render_template(
        'index.html', 
        original_text=original_text, 
        sanitized_text=sanitized_text,
        detected_items=detected_items,
        options=options,
        stats=stats,
        download_id=download_id,
        filename=filename,
        cache_id=cache_id
    )

if __name__ == '__main__':
    print("Starting Flask server...")
    app.run(debug=True, port=5000)